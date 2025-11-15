from google import genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

def create_professional_resume_docx(resume_data, filename="resume.docx"):
    """
    Create a beautifully formatted professional resume DOCX file
    """
    doc = Document()
    
    # Set narrow margins for more space
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # NAME - Large, bold, navy blue
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(resume_data['name'].upper())
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    
    # CONTACT INFO - Centered, smaller
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{resume_data['email']} | {resume_data['phone']} | {resume_data['location']}"
    if resume_data.get('linkedin'):
        contact_text += f" | {resume_data['linkedin']}"
    if resume_data.get('github'):
        contact_text += f" | {resume_data['github']}"
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(64, 64, 64)
    
    # Horizontal line
    doc.add_paragraph('_' * 80)
    
    # PROFESSIONAL SUMMARY
    if resume_data.get('summary'):
        add_section_header(doc, "PROFESSIONAL SUMMARY")
        summary_para = doc.add_paragraph(resume_data['summary'])
        summary_para.paragraph_format.space_after = Pt(12)
        for run in summary_para.runs:
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(64, 64, 64)
    
    # SKILLS
    if resume_data.get('skills'):
        add_section_header(doc, "SKILLS")
        for skill_category, skills_list in resume_data['skills'].items():
            skill_para = doc.add_paragraph(style='List Bullet')
            skill_run = skill_para.add_run(f"{skill_category}: {', '.join(skills_list)}")
            skill_run.font.size = Pt(10.5)
            skill_run.font.color.rgb = RGBColor(64, 64, 64)
        doc.add_paragraph()
    
    # WORK EXPERIENCE
    if resume_data.get('experience'):
        add_section_header(doc, "WORK EXPERIENCE")
        for job in resume_data['experience']:
            # Job title and company - Bold
            job_para = doc.add_paragraph()
            job_title_run = job_para.add_run(f"{job['title']}")
            job_title_run.font.size = Pt(11)
            job_title_run.font.bold = True
            job_title_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Company and duration
            company_para = doc.add_paragraph()
            company_run = company_para.add_run(f"{job['company']} | {job['location']} | {job['duration']}")
            company_run.font.size = Pt(10)
            company_run.font.italic = True
            company_run.font.color.rgb = RGBColor(96, 96, 96)
            
            # Responsibilities
            for responsibility in job['responsibilities']:
                bullet_para = doc.add_paragraph(responsibility, style='List Bullet')
                for run in bullet_para.runs:
                    run.font.size = Pt(10.5)
                    run.font.color.rgb = RGBColor(64, 64, 64)
            
            doc.add_paragraph()  # Space between jobs
    
    # PROJECTS
    if resume_data.get('projects'):
        add_section_header(doc, "PROJECTS")
        for project in resume_data['projects']:
            # Project name - Bold
            project_para = doc.add_paragraph()
            project_name_run = project_para.add_run(project['name'])
            project_name_run.font.size = Pt(11)
            project_name_run.font.bold = True
            project_name_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Technologies
            tech_para = doc.add_paragraph()
            tech_run = tech_para.add_run(f"Technologies: {', '.join(project['technologies'])}")
            tech_run.font.size = Pt(10)
            tech_run.font.italic = True
            tech_run.font.color.rgb = RGBColor(96, 96, 96)
            
            # Description
            desc_para = doc.add_paragraph(f"• {project['description']}")
            for run in desc_para.runs:
                run.font.size = Pt(10.5)
                run.font.color.rgb = RGBColor(64, 64, 64)
            
            doc.add_paragraph()  # Space between projects
    
    # EDUCATION
    if resume_data.get('education'):
        add_section_header(doc, "EDUCATION")
        for edu in resume_data['education']:
            edu_para = doc.add_paragraph()
            degree_run = edu_para.add_run(f"{edu['degree']}")
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True
            degree_run.font.color.rgb = RGBColor(0, 51, 102)
            
            school_para = doc.add_paragraph()
            school_run = school_para.add_run(f"{edu['institution']} | {edu['graduation']} | GPA: {edu.get('gpa', 'N/A')}")
            school_run.font.size = Pt(10)
            school_run.font.italic = True
            school_run.font.color.rgb = RGBColor(96, 96, 96)
            
            doc.add_paragraph()
    
    # CERTIFICATIONS
    if resume_data.get('certifications'):
        add_section_header(doc, "CERTIFICATIONS")
        for cert in resume_data['certifications']:
            cert_para = doc.add_paragraph(style='List Bullet')
            cert_run = cert_para.add_run(f"{cert['name']} - {cert['issuer']} ({cert['year']})")
            cert_run.font.size = Pt(10.5)
            cert_run.font.color.rgb = RGBColor(64, 64, 64)
    
    # Save the document
    doc.save(filename)
    return filename

def add_section_header(doc, text):
    """Add a formatted section header"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    run = para.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Navy blue
    run.font.name = 'Calibri'

def generate_resume_content_with_gemini(profile_data):
    """
    Use Gemini API to expand and professionalize the resume content
    """
    client = genai.Client(api_key="")
    
    # Convert profile to JSON string for the prompt
    profile_json = json.dumps(profile_data, indent=2)
    
    prompt = f"""You are an expert resume writer. Based on the following profile data, create a professional, detailed resume content. 

Expand brief descriptions into professional, impactful statements. Use action verbs and quantify achievements where possible.

Input Profile:
{profile_json}

Generate a JSON response with this EXACT structure:
{{
    "name": "Full Name",
    "email": "email@example.com",
    "phone": "phone number",
    "location": "City, State",
    "linkedin": "LinkedIn URL (if provided)",
    "github": "GitHub URL (if provided)",
    "summary": "A compelling 3-4 sentence professional summary highlighting key skills and experience",
    "skills": {{
        "Programming Languages": ["list", "of", "languages"],
        "Frameworks & Libraries": ["list", "of", "frameworks"],
        "Tools & Technologies": ["list", "of", "tools"]
    }},
    "experience": [
        {{
            "title": "Job Title",
            "company": "Company Name",
            "location": "City, State",
            "duration": "Start Date - End Date",
            "responsibilities": [
                "Detailed achievement-focused bullet point 1",
                "Detailed achievement-focused bullet point 2",
                "Detailed achievement-focused bullet point 3"
            ]
        }}
    ],
    "projects": [
        {{
            "name": "Project Name",
            "technologies": ["tech1", "tech2"],
            "description": "Detailed project description with impact and results"
        }}
    ],
    "education": [
        {{
            "degree": "Degree Name",
            "institution": "University Name",
            "graduation": "Graduation Year",
            "gpa": "GPA if provided"
        }}
    ],
    "certifications": [
        {{
            "name": "Certification Name",
            "issuer": "Issuing Organization",
            "year": "Year"
        }}
    ]
}}

IMPORTANT: Return ONLY valid JSON, no additional text or explanation."""

    print("🤖 Generating professional resume content with Gemini AI...")
    
    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        # Extract JSON from response
        response_text = response.text.strip()
        
        # Remove markdown code blocks if present
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()
        
        # Parse JSON
        resume_data = json.loads(response_text)
        print("✅ Resume content generated successfully!")
        
        return resume_data
        
    except json.JSONDecodeError as e:
        print(f"❌ Error parsing JSON response: {e}")
        print(f"Response text: {response_text[:500]}")
        return None
    except Exception as e:
        print(f"❌ Error generating content: {e}")
        return None

def get_ats_score_and_suggestions(resume_data):
    """
    Get ATS score and suggestions for the generated resume
    """
    client = genai.Client(api_key="AIzaSyDbA6T8H5U0RdNoBFra3dB7bm6AhYgA_zk")
    
    resume_json = json.dumps(resume_data, indent=2)
    
    prompt = f"""You are an ATS (Applicant Tracking System) expert. Analyze this resume and provide:
1. An ATS compatibility score (0-100)
2. Top 3-5 specific suggestions for improvement

RESUME DATA:
{resume_json}

Provide your response in this JSON format:
{{
    "ats_score": <number 0-100>,
    "score_explanation": "Brief explanation of the score",
    "suggestions": [
        "Specific suggestion 1",
        "Specific suggestion 2",
        "Specific suggestion 3"
    ]
}}

Focus on: keyword optimization, formatting issues, missing information, and ATS parsing concerns.
Return ONLY valid JSON."""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        response_text = response.text.strip()
        
        if response_text.startswith("```json"):
            response_text = response_text.replace("```json", "").replace("```", "").strip()
        elif response_text.startswith("```"):
            response_text = response_text.replace("```", "").strip()
        
        ats_result = json.loads(response_text)
        return ats_result
        
    except Exception as e:
        print(f"⚠️  Could not generate ATS score: {e}")
        return None

def main():
    print("="*70)
    print("   PROFESSIONAL RESUME GENERATOR WITH GEMINI AI")
    print("="*70)
    
    # Load profile template
    profile_file = "profile_template.json"
    
    if not os.path.exists(profile_file):
        print(f"❌ File not found: {profile_file}")
        return
    
    with open(profile_file, 'r', encoding='utf-8') as f:
        profile_data = json.load(f)
    
    print(f"✅ Loaded profile: {profile_data.get('personal_info', {}).get('name', 'Unknown')}")
    
    # Generate professional content with Gemini
    resume_data = generate_resume_content_with_gemini(profile_data)
    
    if resume_data is None:
        print("❌ Failed to generate resume content")
        return
    
    # Create DOCX file
    output_file = "resume.docx"
    
    # Handle file conflicts
    counter = 1
    final_filename = output_file
    while os.path.exists(final_filename):
        name, ext = os.path.splitext(output_file)
        final_filename = f"{name}_{counter}{ext}"
        counter += 1
    
    print(f"\n📝 Creating professional DOCX resume...")
    created_file = create_professional_resume_docx(resume_data, final_filename)
    
    print("\n" + "="*70)
    print(f"✅ SUCCESS! Professional resume created!")
    print(f"📂 File: {os.path.abspath(created_file)}")
    print("="*70)
    
    # Get ATS score and suggestions
    print(f"\n🔍 Analyzing ATS compatibility...")
    ats_result = get_ats_score_and_suggestions(resume_data)
    
    if ats_result:
        print("\n" + "="*70)
        print("   ATS SCORE & SUGGESTIONS")
        print("="*70)
        print(f"\n📊 ATS Score: {ats_result.get('ats_score', 'N/A')}/100")
        print(f"   {ats_result.get('score_explanation', '')}")
        
        suggestions = ats_result.get('suggestions', [])
        if suggestions:
            print(f"\n💡 SUGGESTIONS FOR IMPROVEMENT:")
            for i, suggestion in enumerate(suggestions, 1):
                print(f"   {i}. {suggestion}")
        
        print("\n" + "="*70)

if __name__ == "__main__":
    main()
